"""
Document parser for various file formats (PDF, Word, Text) with OCR support.
"""

import os
import uuid
from pathlib import Path
from typing import Optional, List, Dict, Any
import asyncio
import fitz  # PyMuPDF
import pdfplumber
from docx import Document as DocxDocument
import re
import tempfile

# OCR imports
try:
    import pytesseract
    from pdf2image import convert_from_path
    from PIL import Image
    import cv2
    import numpy as np
    
    # Configure Tesseract path for different environments
    import streamlit as st
    
    def get_tesseract_path():
        """Get Tesseract path for different environments."""
        try:
            # Try Streamlit secrets first
            if hasattr(st, 'secrets') and 'TESSERACT_CMD' in st.secrets:
                return st.secrets['TESSERACT_CMD']
        except:
            pass
        
        # Environment variable
        tesseract_path = os.getenv('TESSERACT_CMD')
        if tesseract_path:
            return tesseract_path
        
        # Default paths for different systems
        if os.name == 'nt':  # Windows
            possible_paths = [
                r'C:\Program Files\Tesseract-OCR\tesseract.exe',
                r'C:\Users\Tesseract-OCR\tesseract.exe'
            ]
            for path in possible_paths:
                if os.path.exists(path):
                    return path
        else:  # Linux/Mac
            return '/usr/bin/tesseract'
        
        return 'tesseract'  # Hope it's in PATH
    
    pytesseract.pytesseract.tesseract_cmd = get_tesseract_path()
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

from src.core.models import (
    Document, DocumentType, DocumentSection, DocumentClause, RiskLevel
)
from src.core.config import get_config


class DocumentParser:
    """
    Parser for legal documents in various formats.
    
    Supports:
    - PDF files (with layout preservation)
    - Word documents (.docx, .doc)
    - Plain text files
    """
    
    def __init__(self):
        """Initialize the document parser."""
        self.config = get_config()
        
        # Check OCR availability
        self.ocr_available = OCR_AVAILABLE
        if self.ocr_available:
            # Configure Tesseract OCR
            try:
                # Try to find Tesseract executable
                if os.name == 'nt':  # Windows
                    possible_paths = [
                        r'C:\Program Files\Tesseract-OCR\tesseract.exe',
                        r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
                        r'C:\Users\inand\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'
                    ]
                    for path in possible_paths:
                        if os.path.exists(path):
                            pytesseract.pytesseract.tesseract_cmd = path
                            break
            except:
                self.ocr_available = False
        
        # Common legal document patterns
        self.section_patterns = [
            r'^(\d+\.?\s+[A-Z][^.]*?)\.?\s*$',  # "1. SECTION TITLE"
            r'^([A-Z][A-Z\s]{2,}?)\.?\s*$',      # "SECTION TITLE"
            r'^(Article\s+[IVX\d]+\.?\s+[^.]*?)\.?\s*$',  # "Article I. Title"
            r'^(Section\s+\d+\.?\d*\.?\s+[^.]*?)\.?\s*$', # "Section 1.1 Title"
        ]
        
        self.clause_indicators = [
            r'\b(shall|will|must|may|agrees?|covenant|undertake|represent|warrant)\b',
            r'\b(liability|obligation|duty|responsibility|right|privilege)\b',
            r'\b(payment|compensation|fee|cost|expense|damage)\b',
            r'\b(termination|expiration|breach|default|violation)\b',
        ]
    
    async def parse_document(
        self, 
        file_path: str, 
        document_type: Optional[DocumentType] = None
    ) -> Document:
        """
        Parse a document from file path.
        
        Args:
            file_path: Path to the document file
            document_type: Optional document type hint
            
        Returns:
            Parsed document object
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        # Determine file type and parse accordingly
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.pdf':
            content, metadata = await self._parse_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            content, metadata = await self._parse_word(file_path)
        elif file_extension == '.txt':
            content, metadata = await self._parse_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Create document object
        document = Document(
            id=str(uuid.uuid4()),
            filename=file_path.name,
            file_path=str(file_path),
            document_type=document_type or DocumentType.OTHER,
            content=content,
            metadata=metadata,
            file_size=file_path.stat().st_size,
            page_count=metadata.get('page_count')
        )
        
        # Parse structure
        document.sections = await self._parse_sections(content)
        document.clauses = await self._parse_clauses(content, document.sections)
        
        return document
    
    async def parse_file(
        self, 
        file_path: str, 
        document_type: Optional[DocumentType] = None
    ) -> Document:
        """
        Alias for parse_document to maintain compatibility.
        
        Args:
            file_path: Path to the document file
            document_type: Optional document type hint
            
        Returns:
            Parsed document object
        """
        return await self.parse_document(file_path, document_type)
    
    async def _parse_pdf(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Parse PDF document with OCR fallback."""
        content = ""
        metadata = {}
        
        try:
            # Use PyMuPDF for metadata and basic extraction
            with fitz.open(str(file_path)) as doc:
                metadata = {
                    'page_count': doc.page_count,
                    'title': doc.metadata.get('title', ''),
                    'author': doc.metadata.get('author', ''),
                    'subject': doc.metadata.get('subject', ''),
                    'creator': doc.metadata.get('creator', ''),
                    'producer': doc.metadata.get('producer', ''),
                    'creation_date': doc.metadata.get('creationDate', ''),
                    'modification_date': doc.metadata.get('modDate', '')
                }
            
            # First, try pdfplumber for better text extraction
            text_extracted = False
            with pdfplumber.open(str(file_path)) as pdf:
                pages_text = []
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        pages_text.append(f"[Page {i+1}]\n{page_text}")
                        text_extracted = True
                
                if pages_text:
                    content = "\n\n".join(pages_text)
            
            # If pdfplumber fails, fallback to PyMuPDF
            if not text_extracted:
                with fitz.open(str(file_path)) as doc:
                    for page_num in range(doc.page_count):
                        page = doc[page_num]
                        page_text = page.get_text()
                        if page_text and page_text.strip():
                            content += f"[Page {page_num+1}]\n{page_text}\n\n"
                            text_extracted = True
            
            # If both fail and OCR is available, use OCR
            if not text_extracted and self.ocr_available:
                print("📸 No text found in PDF, attempting OCR extraction...")
                content = await self._extract_text_with_ocr(file_path)
                metadata['ocr_used'] = True
            
            # If still no content, it might be a problematic PDF
            if not content.strip():
                content = f"[Warning: Unable to extract text from PDF. File may be corrupted, password-protected, or contain only images without OCR capability.]"
                metadata['extraction_failed'] = True
        
        except Exception as e:
            # Last resort: try OCR if available
            if self.ocr_available:
                try:
                    content = await self._extract_text_with_ocr(file_path)
                    metadata['ocr_used'] = True
                    metadata['extraction_error'] = str(e)
                except Exception as ocr_error:
                    raise Exception(f"Error parsing PDF (normal extraction failed: {str(e)}, OCR failed: {str(ocr_error)})")
            else:
                raise Exception(f"Error parsing PDF: {str(e)}")
        
        return content, metadata
    
    async def _parse_word(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Parse Word document."""
        try:
            doc = DocxDocument(str(file_path))
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            content = "\n".join(paragraphs)
            
            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'revision': core_props.revision or '',
                'paragraph_count': len(paragraphs)
            }
            
            return content, metadata
            
        except Exception as e:
            raise Exception(f"Error parsing Word document: {str(e)}")
    
    async def _parse_text(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Parse plain text document."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Basic metadata
            metadata = {
                'encoding': 'utf-8',
                'line_count': len(content.splitlines()),
                'char_count': len(content),
                'word_count': len(content.split())
            }
            
            return content, metadata
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    content = file.read()
                metadata = {'encoding': 'latin-1'}
                return content, metadata
            except Exception as e:
                raise Exception(f"Error parsing text file: {str(e)}")
    
    async def _parse_sections(self, content: str) -> List[DocumentSection]:
        """Parse document sections based on common legal document patterns."""
        sections = []
        lines = content.split('\n')
        current_section = None
        section_content = []
        position = 0
        
        for line_num, line in enumerate(lines):
            line = line.strip()
            
            # Check if line is a section header
            is_section_header = False
            section_title = None
            section_number = None
            
            for pattern in self.section_patterns:
                match = re.match(pattern, line, re.IGNORECASE)
                if match:
                    is_section_header = True
                    section_title = match.group(1).strip()
                    
                    # Extract section number if present
                    number_match = re.match(r'^(\d+\.?\d*)', section_title)
                    if number_match:
                        section_number = number_match.group(1)
                    
                    break
            
            if is_section_header and section_title:
                # Save previous section
                if current_section and section_content:
                    current_section.content = '\n'.join(section_content)
                    current_section.end_position = position
                    sections.append(current_section)
                
                # Start new section
                current_section = DocumentSection(
                    id=str(uuid.uuid4()),
                    title=section_title,
                    content="",
                    section_number=section_number,
                    start_position=position,
                    end_position=position,
                    page_number=self._estimate_page_number(line_num, len(lines))
                )
                section_content = []
            
            elif current_section:
                section_content.append(line)
            
            position += len(line) + 1  # +1 for newline
        
        # Save last section
        if current_section and section_content:
            current_section.content = '\n'.join(section_content)
            current_section.end_position = position
            sections.append(current_section)
        
        # If no sections found, create a single section with all content
        if not sections:
            sections.append(DocumentSection(
                id=str(uuid.uuid4()),
                title="Document Content",
                content=content,
                start_position=0,
                end_position=len(content)
            ))
        
        return sections
    
    async def _parse_clauses(
        self, 
        content: str, 
        sections: List[DocumentSection]
    ) -> List[DocumentClause]:
        """Parse individual clauses from the document."""
        clauses = []
        
        for section in sections:
            section_clauses = await self._extract_clauses_from_section(section)
            clauses.extend(section_clauses)
        
        return clauses
    
    async def _extract_clauses_from_section(
        self, 
        section: DocumentSection
    ) -> List[DocumentClause]:
        """Extract clauses from a specific section."""
        clauses = []
        
        # Split section into sentences/clauses
        sentences = re.split(r'[.!?]+', section.content)
        position = section.start_position
        
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) < 20:  # Skip very short sentences
                position += len(sentence) + 1
                continue
            
            # Check if sentence contains legal language
            clause_type = self._identify_clause_type(sentence)
            importance = self._assess_clause_importance(sentence)
            
            if clause_type:
                clause = DocumentClause(
                    id=str(uuid.uuid4()),
                    content=sentence,
                    clause_type=clause_type,
                    section_id=section.id,
                    importance=importance,
                    start_position=position,
                    end_position=position + len(sentence)
                )
                clauses.append(clause)
            
            position += len(sentence) + 1
        
        return clauses
    
    def _identify_clause_type(self, text: str) -> Optional[str]:
        """Identify the type of legal clause."""
        text_lower = text.lower()
        
        # Common clause type patterns
        if any(word in text_lower for word in ['payment', 'pay', 'compensation', 'fee']):
            return 'payment'
        elif any(word in text_lower for word in ['liability', 'liable', 'damages', 'loss']):
            return 'liability'
        elif any(word in text_lower for word in ['termination', 'terminate', 'expire', 'end']):
            return 'termination'
        elif any(word in text_lower for word in ['confidential', 'non-disclosure', 'proprietary']):
            return 'confidentiality'
        elif any(word in text_lower for word in ['intellectual property', 'copyright', 'patent', 'trademark']):
            return 'intellectual_property'
        elif any(word in text_lower for word in ['governing law', 'jurisdiction', 'court', 'dispute']):
            return 'governing_law'
        elif any(word in text_lower for word in ['indemnify', 'indemnification', 'hold harmless']):
            return 'indemnification'
        elif any(word in text_lower for word in ['represent', 'warrant', 'representation', 'warranty']):
            return 'representations_warranties'
        elif any(word in text_lower for word in ['force majeure', 'act of god', 'unforeseeable']):
            return 'force_majeure'
        else:
            # Check for general legal indicators
            for pattern in self.clause_indicators:
                if re.search(pattern, text_lower):
                    return 'general_legal'
        
        return None
    
    def _assess_clause_importance(self, text: str) -> RiskLevel:
        """Assess the importance/risk level of a clause."""
        text_lower = text.lower()
        
        # High importance indicators
        high_risk_terms = [
            'unlimited liability', 'personal guarantee', 'liquidated damages',
            'specific performance', 'criminal liability', 'regulatory violation',
            'material breach', 'immediate termination'
        ]
        
        # Medium importance indicators
        medium_risk_terms = [
            'liability', 'damages', 'breach', 'default', 'termination',
            'indemnification', 'governing law', 'dispute resolution'
        ]
        
        # Low importance indicators
        low_risk_terms = [
            'notice', 'communication', 'interpretation', 'headings',
            'entire agreement', 'amendment', 'severability'
        ]
        
        if any(term in text_lower for term in high_risk_terms):
            return RiskLevel.HIGH
        elif any(term in text_lower for term in medium_risk_terms):
            return RiskLevel.MEDIUM
        elif any(term in text_lower for term in low_risk_terms):
            return RiskLevel.LOW
        else:
            return RiskLevel.MINIMAL
    
    def _estimate_page_number(self, line_num: int, total_lines: int) -> Optional[int]:
        """Estimate page number based on line number (rough approximation)."""
        if total_lines == 0:
            return None
        
        # Assume approximately 50 lines per page
        lines_per_page = 50
        return (line_num // lines_per_page) + 1
    
    async def extract_tables(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract tables from PDF documents."""
        tables = []
        
        if not file_path.endswith('.pdf'):
            return tables
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_tables = page.extract_tables()
                    for i, table in enumerate(page_tables):
                        tables.append({
                            'page': page_num + 1,
                            'table_index': i,
                            'data': table,
                            'bbox': getattr(table, 'bbox', None)
                        })
        except Exception as e:
            print(f"Error extracting tables: {str(e)}")
        
        return tables
    
    async def extract_signatures(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract signature blocks and signing information."""
        signatures = []
        
        try:
            if file_path.endswith('.pdf'):
                with fitz.open(file_path) as doc:
                    for page_num in range(doc.page_count):
                        page = doc[page_num]
                        text = page.get_text()
                        
                        # Look for signature patterns
                        signature_patterns = [
                            r'(?i)(signature|signed|by:)\s*[_\s]*\n([^\n]+)',
                            r'(?i)(name|print name):\s*([^\n]+)',
                            r'(?i)(title):\s*([^\n]+)',
                            r'(?i)(date):\s*([^\n]+)',
                        ]
                        
                        for pattern in signature_patterns:
                            matches = re.finditer(pattern, text)
                            for match in matches:
                                signatures.append({
                                    'page': page_num + 1,
                                    'type': match.group(1).lower(),
                                    'value': match.group(2).strip(),
                                    'position': match.span()
                                })
        
        except Exception as e:
            print(f"Error extracting signatures: {str(e)}")
        
        return signatures

    async def _extract_text_with_ocr(self, file_path: Path) -> str:
        """Extract text using OCR for scanned PDFs or images."""
        if not self.ocr_available:
            return "[OCR not available - install pytesseract, pdf2image, and Tesseract]"
        
        try:
            # Convert PDF to images
            with tempfile.TemporaryDirectory() as temp_dir:
                images = convert_from_path(
                    str(file_path),
                    output_folder=temp_dir,
                    dpi=300,  # Higher DPI for better OCR accuracy
                    fmt='PNG'
                )
                
                extracted_text = []
                
                for i, image in enumerate(images):
                    # Preprocess image for better OCR
                    image_array = np.array(image)
                    
                    # Convert to grayscale if needed
                    if len(image_array.shape) == 3:
                        gray = cv2.cvtColor(image_array, cv2.COLOR_RGB2GRAY)
                    else:
                        gray = image_array
                    
                    # Apply image preprocessing for better OCR
                    # Noise reduction
                    denoised = cv2.medianBlur(gray, 3)
                    
                    # Contrast enhancement
                    enhanced = cv2.convertScaleAbs(denoised, alpha=1.2, beta=10)
                    
                    # Convert back to PIL Image
                    processed_image = Image.fromarray(enhanced)
                    
                    # OCR configuration for better legal document processing
                    custom_config = r'--oem 3 --psm 6 -c tessedit_char_whitelist=ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789.,;:!?()[]{}"-\'$%&*+/=<>@#^_|~` '
                    
                    # Extract text
                    page_text = pytesseract.image_to_string(
                        processed_image,
                        config=custom_config,
                        lang='eng'
                    )
                    
                    if page_text.strip():
                        extracted_text.append(f"[Page {i+1} - OCR]\n{page_text}")
                
                return "\n\n".join(extracted_text) if extracted_text else "[No text extracted via OCR]"
                
        except Exception as e:
            return f"[OCR extraction failed: {str(e)}]"

    def _preprocess_image_for_ocr(self, image: Image.Image) -> Image.Image:
        """Apply image preprocessing techniques for better OCR accuracy."""
        if not self.ocr_available:
            return image
        
        try:
            # Convert to numpy array
            img_array = np.array(image)
            
            # Convert to grayscale if colored
            if len(img_array.shape) == 3:
                gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            else:
                gray = img_array
            
            # Apply Gaussian blur to reduce noise
            blurred = cv2.GaussianBlur(gray, (3, 3), 0)
            
            # Apply threshold to get binary image
            _, binary = cv2.threshold(blurred, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            
            # Morphological operations to clean up the image
            kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2))
            cleaned = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel)
            
            # Convert back to PIL Image
            return Image.fromarray(cleaned)
            
        except Exception as e:
            print(f"Image preprocessing failed: {e}")
            return image
